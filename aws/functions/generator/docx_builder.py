"""
Builds an ATS-friendly DOCX resume from structured JSON data.

Design principles for ATS compatibility:
- Standard fonts (Calibri)
- Simple single-column layout
- No tables, text boxes, images, or headers/footers
- Clear section headings
- Standard bullet characters
"""

from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = "Calibri"
    p.paragraph_format.keep_with_next = True

    border_p = doc.add_paragraph()
    border_p.space_before = Pt(0)
    border_p.space_after = Pt(4)
    border_run = border_p.add_run("_" * 80)
    border_run.font.size = Pt(1)
    border_run.font.color.rgb = RGBColor(100, 100, 100)


def _add_normal(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def build_docx(data: dict[str, Any]) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # --- Header / Contact ---
    if data.get("name"):
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(data["name"])
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_run.font.name = "Calibri"

    contact_parts = []
    if data.get("phone"):
        contact_parts.append(data["phone"])
    if data.get("email"):
        contact_parts.append(data["email"])
    if data.get("linkedin"):
        contact_parts.append(data["linkedin"])
    if data.get("github"):
        contact_parts.append(data["github"])

    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_p.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.name = "Calibri"
        contact_run.font.color.rgb = RGBColor(80, 80, 80)

    # --- Summary ---
    if data.get("summary"):
        _add_heading(doc, "Summary")
        _add_normal(doc, data["summary"])

    # --- Education ---
    if data.get("education"):
        _add_heading(doc, "Education")
        for edu in data["education"]:
            line = ""
            if edu.get("institution"):
                line = edu["institution"]
            if edu.get("dates"):
                line += f"  |  {edu['dates']}" if line else edu["dates"]
            if line:
                _add_normal(doc, line, bold=True)

            if edu.get("degree"):
                _add_normal(doc, edu["degree"], italic=True)

            for detail in edu.get("details", []):
                _add_bullet(doc, detail)

    # --- Experience ---
    if data.get("experience"):
        _add_heading(doc, "Experience")
        for exp in data["experience"]:
            title_line = ""
            if exp.get("title"):
                title_line = exp["title"]
            if exp.get("dates"):
                title_line += f"  |  {exp['dates']}" if title_line else exp["dates"]
            if title_line:
                _add_normal(doc, title_line, bold=True)

            if exp.get("company"):
                _add_normal(doc, exp["company"], italic=True)

            for bullet in exp.get("bullets", []):
                _add_bullet(doc, bullet)

    # --- Projects ---
    if data.get("projects"):
        _add_heading(doc, "Projects")
        for proj in data["projects"]:
            proj_line = proj.get("name", "")
            if proj.get("technologies"):
                proj_line += f"  |  {proj['technologies']}"
            if proj_line:
                _add_normal(doc, proj_line, bold=True)
            if proj.get("description"):
                _add_normal(doc, proj["description"])

    # --- Skills ---
    if data.get("skills"):
        skills = data["skills"]
        has_skills = any(skills.get(k) for k in ["languages", "frameworks", "tools", "other"])
        if has_skills:
            _add_heading(doc, "Technical Skills")
            if skills.get("languages"):
                _add_normal(doc, f"Languages: {', '.join(skills['languages'])}")
            if skills.get("frameworks"):
                _add_normal(doc, f"Frameworks: {', '.join(skills['frameworks'])}")
            if skills.get("tools"):
                _add_normal(doc, f"Tools: {', '.join(skills['tools'])}")
            if skills.get("other"):
                _add_normal(doc, f"Skills: {', '.join(skills['other'])}")

    # --- Certifications ---
    if data.get("certifications"):
        _add_heading(doc, "Certifications")
        for cert in data["certifications"]:
            _add_bullet(doc, cert)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
